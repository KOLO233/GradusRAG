"""Word 文档加载器。

使用 python-docx 解析 .docx 文件，提取段落文本和表格数据。
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path
from typing import List

from src.core.types import Document
from src.ingestion.loaders.base_loader import BaseLoader

logger = logging.getLogger(__name__)


class WordLoader(BaseLoader):
    """Word 文档加载器。

    支持 .docx 格式，提取段落和表格内容。

    Example:
        >>> loader = WordLoader()
        >>> docs = loader.load("data/教案.docx")
    """

    def load(self, file_path: str | Path) -> List[Document]:
        file_path = Path(file_path)
        self._validate_file(file_path)

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required. Install: pip install python-docx"
            )

        logger.info(f"Loading Word: {file_path.name}")

        doc = DocxDocument(str(file_path))
        file_hash = self._compute_file_hash(file_path)
        filename = file_path.name

        # 按一级标题切分章节
        sections = self._extract_sections(doc)

        documents: List[Document] = []
        for i, (title, content) in enumerate(sections):
            if not content.strip():
                continue
            documents.append(Document(
                id=f"{file_hash}_s{i + 1}",
                text=f"## {title}\n\n{content}" if title else content,
                metadata={
                    "source_path": str(file_path),
                    "filename": filename,
                    "doc_type": "word",
                    "doc_id": file_hash,
                    "page": i + 1,
                    "title": title,
                },
            ))

        logger.info(f"  Loaded {len(documents)} sections from {filename}")
        return documents

    @staticmethod
    def _extract_sections(doc) -> List[tuple]:
        """按标题样式切分 Word 文档为章节。"""
        sections = []
        current_title = ""
        current_lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测标题样式
            if para.style and para.style.name and "Heading" in para.style.name:
                if current_lines:
                    sections.append((current_title, "\n".join(current_lines)))
                current_title = text
                current_lines = []
            else:
                current_lines.append(text)

        # 提取表格内容
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(cells))
            if table_text:
                current_lines.append("\n".join(table_text))

        if current_lines:
            sections.append((current_title, "\n".join(current_lines)))

        return sections if sections else [("", "\n".join(
            p.text for p in doc.paragraphs if p.text.strip()
        ))]

    @staticmethod
    def _compute_file_hash(file_path: Path) -> str:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()[:16]
